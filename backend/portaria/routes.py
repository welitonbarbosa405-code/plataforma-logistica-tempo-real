import hashlib
import sqlite3
from datetime import datetime
from io import BytesIO
from pathlib import Path

from flask import Blueprint, jsonify, request, send_file

from database import get_db

portaria_bp = Blueprint('portaria', __name__)

IMAGES_DIR = (Path(__file__).resolve().parent.parent.parent / 'img').resolve()


# ========== CADASTRO ==========

@portaria_bp.route('/api/cadastros', methods=['GET'])
def listar_cadastros():
    try:
        conn = get_db()
        rows = conn.execute('SELECT * FROM cadastro ORDER BY nome').fetchall()
        conn.close()
        return jsonify([dict(r) for r in rows]), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@portaria_bp.route('/api/cadastros/<cpf>', methods=['GET'])
def buscar_cadastro(cpf):
    try:
        conn = get_db()
        row  = conn.execute('SELECT * FROM cadastro WHERE cpf = ?', (cpf,)).fetchone()
        conn.close()
        if not row:
            return jsonify({'error': 'Cadastro não encontrado'}), 404
        return jsonify(dict(row)), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@portaria_bp.route('/api/cadastros', methods=['POST'])
def criar_cadastro():
    try:
        data = request.get_json()
        cpf  = data.get('cpf')
        nome = data.get('nome')
        if not cpf or not nome:
            return jsonify({'error': 'CPF e Nome são obrigatórios'}), 400
        conn   = get_db()
        cursor = conn.cursor()
        if cursor.execute('SELECT cpf FROM cadastro WHERE cpf = ?', (cpf,)).fetchone():
            conn.close()
            return jsonify({'error': 'CPF já cadastrado.'}), 400
        cursor.execute(
            'INSERT INTO cadastro (cpf, nome, telefone, integracao, empresa, transportadora, placa) VALUES (?,?,?,?,?,?,?)',
            (cpf, nome, data.get('telefone'), data.get('integracao'),
             data.get('empresa'), data.get('transportadora'), data.get('placa'))
        )
        conn.commit()
        conn.close()
        return jsonify({'message': 'Cadastro criado com sucesso', 'cpf': cpf}), 201
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@portaria_bp.route('/api/cadastros/<cpf>', methods=['PUT'])
def atualizar_cadastro(cpf):
    try:
        data   = request.get_json()
        conn   = get_db()
        cursor = conn.cursor()
        cursor.execute(
            'UPDATE cadastro SET nome=?, telefone=?, integracao=?, empresa=?, transportadora=?, placa=? WHERE cpf=?',
            (data.get('nome'), data.get('telefone'), data.get('integracao'),
             data.get('empresa'), data.get('transportadora'), data.get('placa'), cpf)
        )
        if cursor.rowcount == 0:
            conn.close()
            return jsonify({'error': 'Cadastro não encontrado'}), 404
        conn.commit()
        conn.close()
        return jsonify({'message': 'Cadastro atualizado com sucesso'}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@portaria_bp.route('/api/cadastros/<cpf>', methods=['DELETE'])
def deletar_cadastro(cpf):
    try:
        conn   = get_db()
        cursor = conn.cursor()
        cursor.execute('DELETE FROM cadastro WHERE cpf = ?', (cpf,))
        if cursor.rowcount == 0:
            conn.close()
            return jsonify({'error': 'Cadastro não encontrado'}), 404
        conn.commit()
        conn.close()
        return jsonify({'message': 'Cadastro deletado com sucesso'}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ========== ENTRADAS ==========

@portaria_bp.route('/api/entradas', methods=['GET'])
def listar_entradas():
    try:
        conn = get_db()
        rows = conn.execute(
            'SELECT * FROM cadastro_entrada ORDER BY data_entrada DESC, hora_entrada DESC'
        ).fetchall()
        conn.close()
        return jsonify([dict(r) for r in rows]), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@portaria_bp.route('/api/entradas/<int:id>', methods=['GET'])
def buscar_entrada(id):
    try:
        conn = get_db()
        row  = conn.execute('SELECT * FROM cadastro_entrada WHERE id = ?', (id,)).fetchone()
        conn.close()
        if not row:
            return jsonify({'error': 'Entrada não encontrada'}), 404
        return jsonify(dict(row)), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@portaria_bp.route('/api/entradas', methods=['POST'])
def criar_entrada():
    try:
        data = request.get_json()
        cpf  = data.get('cpf')
        nome = data.get('nome')
        if not cpf or not nome:
            return jsonify({'error': 'CPF e Nome são obrigatórios'}), 400
        agora       = datetime.now()
        data_entrada = data.get('data') or agora.strftime('%Y-%m-%d')
        hora_entrada = agora.strftime('%H:%M')
        conn   = get_db()
        cursor = conn.cursor()
        cursor.execute('''
            INSERT INTO cadastro_entrada
            (data, tipo, cpf, nome, telefone, integracao, placa, transportadora,
             numero_coleta, empresa, solicitante, vigilante, data_entrada, hora_entrada)
            VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?)
        ''', (data.get('data'), data.get('tipo'), cpf, nome, data.get('telefone'),
              data.get('integracao'), data.get('placa'), data.get('transportadora'),
              data.get('numero_coleta'), data.get('empresa'), data.get('solicitante'),
              data.get('vigilante'), data_entrada, hora_entrada))
        entrada_id = cursor.lastrowid
        conn.commit()
        conn.close()
        return jsonify({'message': 'Entrada registrada com sucesso', 'id': entrada_id}), 201
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@portaria_bp.route('/api/entradas/<int:id>/saida', methods=['PUT'])
def registrar_saida(id):
    try:
        agora  = datetime.now()
        conn   = get_db()
        cursor = conn.cursor()
        cursor.execute(
            'UPDATE cadastro_entrada SET data_saida=?, hora_saida=? WHERE id=?',
            (agora.strftime('%Y-%m-%d'), agora.strftime('%H:%M'), id)
        )
        if cursor.rowcount == 0:
            conn.close()
            return jsonify({'error': 'Entrada não encontrada'}), 404
        conn.commit()
        conn.close()
        return jsonify({'message': 'Saída registrada com sucesso'}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@portaria_bp.route('/api/entradas/<int:id>', methods=['PUT'])
def atualizar_entrada(id):
    try:
        data   = request.get_json()
        conn   = get_db()
        cursor = conn.cursor()
        cursor.execute('''
            UPDATE cadastro_entrada SET
            data=?, tipo=?, cpf=?, nome=?, telefone=?, integracao=?, placa=?,
            transportadora=?, numero_coleta=?, empresa=?, solicitante=?, vigilante=?,
            data_entrada=?, hora_entrada=?, data_saida=?, hora_saida=?
            WHERE id=?
        ''', (data.get('data'), data.get('tipo'), data.get('cpf'), data.get('nome'),
              data.get('telefone'), data.get('integracao'), data.get('placa'),
              data.get('transportadora'), data.get('numero_coleta'), data.get('empresa'),
              data.get('solicitante'), data.get('vigilante'), data.get('data_entrada'),
              data.get('hora_entrada'), data.get('data_saida'), data.get('hora_saida'), id))
        if cursor.rowcount == 0:
            conn.close()
            return jsonify({'error': 'Entrada não encontrada'}), 404
        conn.commit()
        conn.close()
        return jsonify({'message': 'Entrada atualizada com sucesso'}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@portaria_bp.route('/api/entradas/<int:id>', methods=['DELETE'])
def deletar_entrada(id):
    try:
        conn   = get_db()
        cursor = conn.cursor()
        cursor.execute('DELETE FROM cadastro_entrada WHERE id = ?', (id,))
        if cursor.rowcount == 0:
            conn.close()
            return jsonify({'error': 'Entrada não encontrada'}), 404
        conn.commit()
        conn.close()
        return jsonify({'message': 'Entrada deletada com sucesso'}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ========== TICKET WORD ==========

@portaria_bp.route('/api/entradas/<int:id>/ticket', methods=['GET'])
def gerar_ticket_entrada(id):
    try:
        from docx import Document
        from docx.shared import Inches, Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT

        conn   = get_db()
        row    = conn.execute('SELECT * FROM cadastro_entrada WHERE id = ?', (id,)).fetchone()
        conn.close()
        if not row:
            return jsonify({'error': 'Entrada não encontrada'}), 404
        entrada = dict(row)

        doc = Document()
        for section in doc.sections:
            section.top_margin    = Cm(1)
            section.bottom_margin = Cm(1)
            section.left_margin   = Cm(1.5)
            section.right_margin  = Cm(1.5)

        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        p_titulo = table.cell(0, 0).paragraphs[0]
        run = p_titulo.add_run('TICKET DE ENTRADA')
        run.bold = True
        run.font.size = Pt(18)
        p_titulo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_logo = table.cell(0, 1).paragraphs[0]
        p_logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        logo_path = IMAGES_DIR / 'logo_kuhn.png'
        if logo_path.exists():
            p_logo.add_run().add_picture(str(logo_path), width=Inches(1))

        doc.add_paragraph()
        p_num = doc.add_paragraph()
        p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_num.add_run('Nº FORMULÁRIO: ').font.size = Pt(10)
        run_num = p_num.add_run(f'{entrada["id"]}')
        run_num.bold = True
        run_num.font.size = Pt(14)

        doc.add_paragraph('─' * 60)
        p_tipo = doc.add_paragraph()
        p_tipo.add_run('TIPO ENTRADA: ').font.size = Pt(10)
        run_tipo = p_tipo.add_run(f'{entrada["tipo"] or "ENTREGA"}')
        run_tipo.bold = True
        run_tipo.font.size = Pt(12)

        doc.add_paragraph('─' * 60)
        p_nome_lbl = doc.add_paragraph()
        p_nome_lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_nome_lbl.add_run('NOME').font.size = Pt(9)
        p_nome = doc.add_paragraph()
        p_nome.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_nome = p_nome.add_run(f'{entrada["nome"] or "-"}')
        run_nome.bold = True
        run_nome.font.size = Pt(14)

        doc.add_paragraph('─' * 60)
        try:
            data_fmt = datetime.strptime(entrada.get('data_entrada', ''), '%Y-%m-%d').strftime('%d/%m/%Y')
        except Exception:
            data_fmt = entrada.get('data_entrada') or '-'

        tbl = doc.add_table(rows=2, cols=3)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(['PLACA', 'DATA', 'HORÁRIO']):
            p = tbl.cell(0, i).paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(h).font.size = Pt(9)
        for i, v in enumerate([entrada.get('placa') or '-', data_fmt, entrada.get('hora_entrada') or '-']):
            p = tbl.cell(1, i).paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(v))
            run.bold = True
            run.font.size = Pt(12)

        doc.add_paragraph('─' * 60)
        p_vig = doc.add_paragraph()
        p_vig.add_run('VIGILANTE: ').font.size = Pt(9)
        p_vig.add_run(f'{entrada.get("vigilante") or "________________"}').font.size = Pt(11)
        p_sol = doc.add_paragraph()
        p_sol.add_run('SOLICITANTE: ').font.size = Pt(9)
        p_sol.add_run(f'{entrada.get("solicitante") or "________________"}').font.size = Pt(11)

        stream = BytesIO()
        doc.save(stream)
        stream.seek(0)
        nome_arquivo = f'ticket_entrada_{entrada["id"]}_{(entrada["nome"] or "motorista").replace(" ","_")}.docx'
        return send_file(stream, as_attachment=True, download_name=nome_arquivo,
                         mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ========== AUTENTICAÇÃO ==========

@portaria_bp.route('/api/auth/login', methods=['POST'])
def login():
    try:
        data  = request.get_json()
        email = data.get('email')
        senha = data.get('senha')
        if not email or not senha:
            return jsonify({'error': 'Email e senha são obrigatórios'}), 400
        conn    = get_db()
        usuario = conn.execute('SELECT * FROM usuarios WHERE email = ?', (email,)).fetchone()
        conn.close()
        if not usuario:
            return jsonify({'error': 'Email ou senha incorretos'}), 401
        if usuario['senha_hash'] != hashlib.sha256(senha.encode()).hexdigest():
            return jsonify({'error': 'Email ou senha incorretos'}), 401
        return jsonify({'message': 'Login realizado com sucesso', 'usuario': {
            'id': usuario['id'], 'nome': usuario['nome'],
            'email': usuario['email'], 'perfil': usuario['perfil'],
            'criado_em': usuario['criado_em']
        }}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@portaria_bp.route('/api/auth/verificar', methods=['GET'])
def verificar_sessao():
    return jsonify({'message': 'Sessão válida'}), 200


# ========== JANELA ==========

@portaria_bp.route('/api/janela', methods=['GET'])
def get_janelas():
    try:
        conn  = get_db()
        rows  = conn.execute('SELECT * FROM janela ORDER BY dia_semana, horario, transportadora').fetchall()
        conn.close()
        return jsonify([dict(r) for r in rows]), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@portaria_bp.route('/api/janela/<int:janela_id>', methods=['GET'])
def get_janela(janela_id):
    try:
        conn = get_db()
        row  = conn.execute('SELECT * FROM janela WHERE id = ?', (janela_id,)).fetchone()
        conn.close()
        if row:
            return jsonify(dict(row)), 200
        return jsonify({'error': 'Janela não encontrada'}), 404
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@portaria_bp.route('/api/janela', methods=['POST'])
def criar_janela():
    try:
        data           = request.get_json()
        filial         = data.get('filial')
        transportadora = data.get('transportadora')
        dia_semana     = data.get('dia_semana')
        horario        = data.get('horario')
        tipo           = data.get('tipo', 'Coleta')
        if not all([filial, transportadora, dia_semana, horario]):
            return jsonify({'error': 'Filial, Transportadora, Dia da semana e Horário são obrigatórios'}), 400
        conn   = get_db()
        cursor = conn.cursor()
        if cursor.execute(
            'SELECT id FROM janela WHERE transportadora=? AND dia_semana=? AND horario=?',
            (transportadora, dia_semana, horario)
        ).fetchone():
            conn.close()
            return jsonify({'error': 'Esta transportadora já possui horário cadastrado neste dia e horário'}), 400
        cursor.execute(
            'INSERT INTO janela (filial,transportadora,dia_semana,horario,tipo) VALUES (?,?,?,?,?)',
            (filial, transportadora, dia_semana, horario, tipo)
        )
        conn.commit()
        janela_id = cursor.lastrowid
        conn.close()
        return jsonify({'message': 'Janela cadastrada com sucesso', 'id': janela_id}), 201
    except sqlite3.IntegrityError:
        return jsonify({'error': 'Esta transportadora já possui horário cadastrado neste dia e horário'}), 400
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@portaria_bp.route('/api/janela/<int:janela_id>', methods=['PUT'])
def atualizar_janela(janela_id):
    try:
        data           = request.get_json()
        filial         = data.get('filial')
        transportadora = data.get('transportadora')
        dia_semana     = data.get('dia_semana')
        horario        = data.get('horario')
        tipo           = data.get('tipo')
        if not all([filial, transportadora, dia_semana, horario]):
            return jsonify({'error': 'Campos obrigatórios não preenchidos'}), 400
        conn   = get_db()
        cursor = conn.cursor()
        if cursor.execute(
            'SELECT id FROM janela WHERE transportadora=? AND dia_semana=? AND horario=? AND id!=?',
            (transportadora, dia_semana, horario, janela_id)
        ).fetchone():
            conn.close()
            return jsonify({'error': 'Esta transportadora já possui horário cadastrado neste dia e horário'}), 400
        cursor.execute(
            'UPDATE janela SET filial=?,transportadora=?,dia_semana=?,horario=?,tipo=? WHERE id=?',
            (filial, transportadora, dia_semana, horario, tipo, janela_id)
        )
        conn.commit()
        conn.close()
        return jsonify({'message': 'Janela atualizada com sucesso'}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@portaria_bp.route('/api/janela/<int:janela_id>', methods=['DELETE'])
def deletar_janela(janela_id):
    try:
        conn = get_db()
        conn.execute('DELETE FROM janela WHERE id = ?', (janela_id,))
        conn.commit()
        conn.close()
        return jsonify({'message': 'Janela deletada com sucesso'}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ========== DASHBOARD ==========

@portaria_bp.route('/api/dashboard/metrics', methods=['GET'])
def dashboard_metrics():
    try:
        conn = get_db()
        total_entradas       = conn.execute('SELECT COUNT(*) FROM cadastro_entrada').fetchone()[0]
        tempo_medio_result   = conn.execute('''
            SELECT AVG((julianday(data_saida||' '||hora_saida)-julianday(data_entrada||' '||hora_entrada))*24*60)
            FROM cadastro_entrada WHERE data_saida IS NOT NULL AND hora_saida IS NOT NULL
        ''').fetchone()[0]
        total_transportadoras = conn.execute('''
            SELECT COUNT(DISTINCT transportadora) FROM cadastro_entrada
            WHERE transportadora IS NOT NULL AND transportadora != ""
        ''').fetchone()[0]
        total_janelas = conn.execute('SELECT COUNT(*) FROM janela').fetchone()[0]
        conn.close()
        return jsonify({
            'total_entradas': total_entradas,
            'tempo_medio_permanencia': int(tempo_medio_result) if tempo_medio_result else 0,
            'total_transportadoras': total_transportadoras,
            'total_janelas': total_janelas
        }), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@portaria_bp.route('/api/dashboard/entradas-horario', methods=['GET'])
def dashboard_entradas_horario():
    try:
        conn = get_db()
        rows = conn.execute('''
            SELECT CAST(substr(hora_entrada,1,2) AS INTEGER) as hora, COUNT(*) as total
            FROM cadastro_entrada WHERE hora_entrada IS NOT NULL GROUP BY hora ORDER BY hora
        ''').fetchall()
        conn.close()
        return jsonify({'horas': [f"{h:02d}:00" for h, _ in rows], 'quantidades': [t for _, t in rows]}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@portaria_bp.route('/api/dashboard/entradas-transportadora', methods=['GET'])
def dashboard_entradas_transportadora():
    try:
        conn = get_db()
        rows = conn.execute('''
            SELECT transportadora, COUNT(*) as total FROM cadastro_entrada
            WHERE transportadora IS NOT NULL AND transportadora != ""
            GROUP BY transportadora ORDER BY total DESC LIMIT 10
        ''').fetchall()
        conn.close()
        return jsonify({'transportadoras': [r[0] for r in rows], 'quantidades': [r[1] for r in rows]}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@portaria_bp.route('/api/dashboard/entradas-dia-semana', methods=['GET'])
def dashboard_entradas_dia_semana():
    try:
        conn = get_db()
        rows = conn.execute('''
            SELECT strftime('%w', data_entrada) as dia_num, COUNT(*) as total
            FROM cadastro_entrada WHERE data_entrada IS NOT NULL GROUP BY dia_num ORDER BY dia_num
        ''').fetchall()
        conn.close()
        q = [0] * 7
        for dia_num, total in rows:
            q[int(dia_num)] = total
        dias = ['Segunda-Feira','Terça-Feira','Quarta-Feira','Quinta-Feira','Sexta-Feira','Sábado','Domingo']
        return jsonify({'dias': dias, 'quantidades': [q[1],q[2],q[3],q[4],q[5],q[6],q[0]]}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@portaria_bp.route('/api/dashboard/tempo-permanencia', methods=['GET'])
def dashboard_tempo_permanencia():
    try:
        conn = get_db()
        rows = conn.execute('''
            SELECT CASE
                WHEN (julianday(data_saida||' '||hora_saida)-julianday(data_entrada||' '||hora_entrada))*24*60 < 15 THEN 'Rápido (< 15min)'
                WHEN (julianday(data_saida||' '||hora_saida)-julianday(data_entrada||' '||hora_entrada))*24*60 <= 120 THEN 'Normal (15min - 2h)'
                ELSE 'Longo (> 2h)'
            END as categoria, COUNT(*) as total
            FROM cadastro_entrada WHERE data_saida IS NOT NULL AND hora_saida IS NOT NULL GROUP BY categoria
        ''').fetchall()
        conn.close()
        return jsonify({'categorias': [r[0] for r in rows], 'quantidades': [r[1] for r in rows]}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@portaria_bp.route('/api/dashboard/estatisticas-transportadora', methods=['GET'])
def dashboard_estatisticas_transportadora():
    try:
        conn = get_db()
        rows = conn.execute('''
            SELECT e.transportadora, COUNT(DISTINCT e.id) as total_entradas,
                AVG(CASE WHEN e.data_saida IS NOT NULL AND e.hora_saida IS NOT NULL
                    THEN (julianday(e.data_saida||' '||e.hora_saida)-julianday(e.data_entrada||' '||e.hora_entrada))*24*60
                    ELSE NULL END) as tempo_medio,
                SUM(CASE WHEN date(e.data_entrada)=date('now') THEN 1 ELSE 0 END) as entradas_hoje,
                (SELECT COUNT(*) FROM janela j WHERE j.transportadora=e.transportadora) as janelas_cadastradas
            FROM cadastro_entrada e
            WHERE e.transportadora IS NOT NULL AND e.transportadora != ""
            GROUP BY e.transportadora ORDER BY total_entradas DESC
        ''').fetchall()
        conn.close()
        return jsonify([{
            'transportadora': r[0] or 'Não informado',
            'total_entradas': r[1],
            'tempo_medio': int(r[2]) if r[2] else 0,
            'entradas_hoje': r[3],
            'janelas_cadastradas': r[4] or 0
        } for r in rows]), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ========== USUÁRIOS ==========

_PERM_COLS = [
    'perm_cadastro_criar', 'perm_cadastro_editar', 'perm_cadastro_excluir',
    'perm_entrada_criar', 'perm_entrada_editar', 'perm_entrada_excluir', 'perm_entrada_registrar_saida',
    'perm_janela_criar', 'perm_janela_editar', 'perm_janela_excluir',
    'perm_dashboard_visualizar', 'perm_usuarios_gerenciar'
]


def _row_to_usuario(row):
    u = dict(row)
    for k in u:
        if k.startswith('perm_'):
            u[k] = bool(u[k])
    return u


@portaria_bp.route('/api/usuarios', methods=['GET'])
def listar_usuarios():
    try:
        conn = get_db()
        rows = conn.execute(f'''
            SELECT id,nome,email,perfil,criado_em,{",".join(_PERM_COLS)} FROM usuarios ORDER BY nome
        ''').fetchall()
        conn.close()
        return jsonify([_row_to_usuario(r) for r in rows]), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@portaria_bp.route('/api/usuarios/<int:usuario_id>', methods=['GET'])
def buscar_usuario(usuario_id):
    try:
        conn = get_db()
        row  = conn.execute(f'''
            SELECT id,nome,email,perfil,criado_em,{",".join(_PERM_COLS)} FROM usuarios WHERE id=?
        ''', (usuario_id,)).fetchone()
        conn.close()
        if not row:
            return jsonify({'error': 'Usuário não encontrado'}), 404
        return jsonify(_row_to_usuario(row)), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@portaria_bp.route('/api/usuarios', methods=['POST'])
def criar_usuario():
    try:
        data   = request.get_json()
        nome   = data.get('nome')
        email  = data.get('email')
        senha  = data.get('senha')
        perfil = data.get('perfil', 'porteiro')
        if not all([nome, email, senha]):
            return jsonify({'error': 'Nome, Email e Senha são obrigatórios'}), 400
        perfis_validos = ['porteiro', 'supervisor', 'administrador', 'operador', 'visualizador', 'admin']
        if perfil not in perfis_validos:
            return jsonify({'error': 'Perfil inválido'}), 400
        admins = ['administrador', 'admin']
        supvs  = ['administrador', 'supervisor', 'admin']
        perms  = {
            'perm_cadastro_criar':         data.get('perm_cadastro_criar', True),
            'perm_cadastro_editar':        data.get('perm_cadastro_editar', perfil in supvs),
            'perm_cadastro_excluir':       data.get('perm_cadastro_excluir', perfil in admins),
            'perm_entrada_criar':          data.get('perm_entrada_criar', True),
            'perm_entrada_editar':         data.get('perm_entrada_editar', perfil in supvs),
            'perm_entrada_excluir':        data.get('perm_entrada_excluir', perfil in admins),
            'perm_entrada_registrar_saida': data.get('perm_entrada_registrar_saida', True),
            'perm_janela_criar':           data.get('perm_janela_criar', perfil in supvs),
            'perm_janela_editar':          data.get('perm_janela_editar', perfil in supvs),
            'perm_janela_excluir':         data.get('perm_janela_excluir', perfil in admins),
            'perm_dashboard_visualizar':   data.get('perm_dashboard_visualizar', True),
            'perm_usuarios_gerenciar':     data.get('perm_usuarios_gerenciar', perfil in admins),
        }
        senha_hash = hashlib.sha256(senha.encode()).hexdigest()
        conn   = get_db()
        cursor = conn.cursor()
        if cursor.execute('SELECT id FROM usuarios WHERE email=?', (email,)).fetchone():
            conn.close()
            return jsonify({'error': 'Email já cadastrado'}), 400
        cursor.execute(f'''
            INSERT INTO usuarios (nome,email,senha_hash,perfil,{",".join(_PERM_COLS)})
            VALUES (?,?,?,?,{",".join(["?"]*len(_PERM_COLS))})
        ''', (nome, email, senha_hash, perfil, *[int(perms[c]) for c in _PERM_COLS]))
        conn.commit()
        usuario_id = cursor.lastrowid
        conn.close()
        return jsonify({'message': 'Usuário criado com sucesso', 'id': usuario_id}), 201
    except sqlite3.IntegrityError:
        return jsonify({'error': 'Email já cadastrado'}), 400
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@portaria_bp.route('/api/usuarios/<int:usuario_id>', methods=['PUT'])
def atualizar_usuario(usuario_id):
    try:
        data   = request.get_json()
        conn   = get_db()
        cursor = conn.cursor()
        if not cursor.execute('SELECT id FROM usuarios WHERE id=?', (usuario_id,)).fetchone():
            conn.close()
            return jsonify({'error': 'Usuário não encontrado'}), 404
        email = data.get('email')
        if email and cursor.execute('SELECT id FROM usuarios WHERE email=? AND id!=?', (email, usuario_id)).fetchone():
            conn.close()
            return jsonify({'error': 'Email já está em uso por outro usuário'}), 400
        updates, params = [], []
        for field in ['nome', 'email', 'perfil']:
            if data.get(field):
                updates.append(f'{field}=?')
                params.append(data.get(field))
        if data.get('senha'):
            updates.append('senha_hash=?')
            params.append(hashlib.sha256(data['senha'].encode()).hexdigest())
        for campo in _PERM_COLS:
            if campo in data:
                updates.append(f'{campo}=?')
                params.append(int(bool(data[campo])))
        if not updates:
            conn.close()
            return jsonify({'error': 'Nenhum campo para atualizar'}), 400
        params.append(usuario_id)
        cursor.execute(f'UPDATE usuarios SET {",".join(updates)} WHERE id=?', params)
        conn.commit()
        conn.close()
        return jsonify({'message': 'Usuário atualizado com sucesso'}), 200
    except sqlite3.IntegrityError:
        return jsonify({'error': 'Email já está em uso'}), 400
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@portaria_bp.route('/api/usuarios/<int:usuario_id>', methods=['DELETE'])
def deletar_usuario(usuario_id):
    try:
        conn   = get_db()
        cursor = conn.cursor()
        if not cursor.execute('SELECT id FROM usuarios WHERE id=?', (usuario_id,)).fetchone():
            conn.close()
            return jsonify({'error': 'Usuário não encontrado'}), 404
        total_admins  = cursor.execute("SELECT COUNT(*) FROM usuarios WHERE perfil='administrador'").fetchone()[0]
        perfil_usuario = cursor.execute('SELECT perfil FROM usuarios WHERE id=?', (usuario_id,)).fetchone()[0]
        if perfil_usuario == 'administrador' and total_admins <= 1:
            conn.close()
            return jsonify({'error': 'Não é possível deletar o último administrador'}), 400
        cursor.execute('DELETE FROM usuarios WHERE id=?', (usuario_id,))
        conn.commit()
        conn.close()
        return jsonify({'message': 'Usuário deletado com sucesso'}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500
